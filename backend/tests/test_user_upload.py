"""Тесты загрузки источников пользователем.

Файлы собираются в памяти, диск и сеть не используются.
"""

import io

import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.modules.sources.user_upload import (
    UploadError,
    build_source,
    extract_text,
    extract_txt,
    guess_authors,
    guess_doi,
    guess_title,
    guess_year,
    load_uploaded_source,
    merge_with_found,
)
from app.modules.sources.openalex import Source

ARTICLE = (
    "УДК 347.51 ГРНТИ 10.27.65\n"
    "\n"
    "СУБСИДИАРНАЯ ОТВЕТСТВЕННОСТЬ КОНТРОЛИРУЮЩИХ ДОЛЖНИКА ЛИЦ\n"
    "\n"
    "Иванов И. И.\n"
    "Московский государственный университет, 2024\n"
    "\n"
    "Аннотация. В статье рассматриваются условия привлечения "
    "контролирующих должника лиц к субсидиарной ответственности. "
    + "Автор анализирует практику арбитражных судов. " * 30
)


def docx_bytes(paragraphs, table_rows=None):
    import docx
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        t = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestExtractText:
    def test_utf8(self):
        assert "Привет" in extract_txt("Привет мир".encode("utf-8"))

    def test_cp1251_detected(self):
        text = "Субсидиарная ответственность контролирующих лиц"
        assert extract_txt(text.encode("cp1251")) == text

    def test_bom_stripped(self):
        assert extract_txt("\ufeffТекст".encode("utf-8-sig")) == "Текст"

    def test_docx_paragraphs(self):
        data = docx_bytes(["Первый абзац", "Второй абзац"])
        out = extract_text("a.docx", data)
        assert "Первый абзац" in out and "Второй абзац" in out

    def test_docx_tables_included(self):
        data = docx_bytes(["Текст"], [["Год", "Дел"], ["2024", "150"]])
        out = extract_text("a.docx", data)
        assert "Год | Дел" in out

    def test_unsupported_extension(self):
        with pytest.raises(UploadError, match="не поддерживается"):
            extract_text("scan.jpg", b"x" * 1000)

    def test_empty_file(self):
        with pytest.raises(UploadError, match="пустой"):
            extract_text("a.txt", b"")

    def test_too_large(self):
        with pytest.raises(UploadError, match="МБ"):
            extract_text("a.pdf", b"x" * (26 * 1024 * 1024))

    def test_broken_pdf(self):
        with pytest.raises(UploadError, match="не удалось прочитать PDF"):
            extract_text("a.pdf", b"%PDF-1.4 broken garbage")

    def test_legacy_doc_hint(self):
        with pytest.raises(UploadError, match=r"\.docx"):
            extract_text("a.docx", b"\xd0\xcf\x11\xe0old binary doc")

    def test_rtf(self):
        rtf = rb"{\rtf1\ansi\ansicpg1251 \'d2\'e5\'ea\'f1\'f2 \par}"
        assert "Текст" in extract_text("a.rtf", rtf)


class TestMetadataGuess:
    def test_title_skips_udk(self):
        title = guess_title(ARTICLE, "file.pdf")
        assert title == "СУБСИДИАРНАЯ ОТВЕТСТВЕННОСТЬ КОНТРОЛИРУЮЩИХ ДОЛЖНИКА ЛИЦ"

    def test_title_falls_back_to_filename(self):
        assert guess_title("ааа", "моя_статья.pdf") == "моя статья"

    def test_title_skips_doi_line(self):
        text = "DOI: 10.1234/abcd.2024.11\nНастоящий заголовок статьи здесь\n"
        assert guess_title(text, "f.pdf") == "Настоящий заголовок статьи здесь"

    def test_year(self):
        assert guess_year(ARTICLE) == 2024

    def test_year_absent(self):
        assert guess_year("текст без дат " * 50) is None

    def test_authors(self):
        assert "Иванов И. И." in guess_authors(ARTICLE)

    def test_doi(self):
        text = "DOI: 10.31857/S013207690012345-6\n" + ARTICLE
        assert guess_doi(text) == "10.31857/S013207690012345-6"

    def test_doi_absent(self):
        assert guess_doi(ARTICLE) is None


class TestBuildSource:
    def test_builds_with_fulltext(self):
        r = build_source("a.pdf", ARTICLE)
        assert r.source.provider == "user_upload"
        assert r.source.fulltext.startswith("УДК")
        assert r.source.relevance == 1.0
        assert r.source.year == 2024

    def test_scan_without_text_layer_rejected(self):
        with pytest.raises(UploadError, match="скан"):
            build_source("scan.pdf", "три слова тут")

    def test_explicit_title_wins(self):
        r = build_source("a.pdf", ARTICLE, title="Моё название")
        assert r.source.title == "Моё название"

    def test_warns_about_guessed_title(self):
        r = build_source("a.pdf", ARTICLE)
        assert any("автоматически" in w for w in r.warnings)

    def test_warns_about_missing_year(self):
        text = "Заголовок без выходных данных\n" + "слово " * 200
        r = build_source("a.pdf", text)
        assert any("год" in w for w in r.warnings)

    def test_load_from_bytes(self):
        r = load_uploaded_source("a.txt", ARTICLE.encode("cp1251"))
        assert r.chars > 400
        assert r.source.year == 2024


class TestMerge:
    def test_uploaded_go_first(self):
        up = [build_source("a.pdf", ARTICLE).source]
        found = [Source(title=f"Найденная {i}", doi=f"10/{i}")
                 for i in range(6)]
        merged = merge_with_found(up, found, limit=6)
        assert merged[0].provider == "user_upload"
        assert len(merged) == 6

    def test_uploaded_never_dropped(self):
        up = [build_source(f"{i}.pdf", ARTICLE, title=f"Моя {i}").source
              for i in range(8)]
        merged = merge_with_found(up, [Source(title="Найденная")], limit=6)
        assert len(merged) == 8
        assert all(s.provider == "user_upload" for s in merged)

    def test_duplicate_title_not_added_twice(self):
        up = [build_source("a.pdf", ARTICLE, title="Одна статья").source]
        found = [Source(title="Одна статья!", doi="10/1")]
        assert len(merge_with_found(up, found, limit=6)) == 1


class TestUploadAPI:
    @pytest.fixture
    def client(self):
        return TestClient(app)

    def test_upload_source_accepted(self, client):
        r = client.post(
            "/api/v1/projects/upload/sources",
            files=[("files", ("a.txt", ARTICLE.encode("utf-8"), "text/plain"))],
        )
        assert r.status_code == 200
        body = r.json()
        assert len(body["accepted"]) == 1
        assert not body["rejected"]
        assert body["accepted"][0]["source"]["provider"] == "user_upload"

    def test_bad_file_reported_not_raised(self, client):
        r = client.post(
            "/api/v1/projects/upload/sources",
            files=[
                ("files", ("good.txt", ARTICLE.encode(), "text/plain")),
                ("files", ("bad.jpg", b"x" * 900, "image/jpeg")),
            ],
        )
        assert r.status_code == 200
        assert len(r.json()["accepted"]) == 1
        assert len(r.json()["rejected"]) == 1

    def test_manual_title_applied(self, client):
        r = client.post(
            "/api/v1/projects/upload/sources",
            files=[("files", ("a.txt", ARTICLE.encode(), "text/plain"))],
            data={"titles": "Ручное название"},
        )
        assert r.json()["accepted"][0]["source"]["title"] == "Ручное название"

    def test_too_many_files(self, client):
        files = [("files", (f"{i}.txt", ARTICLE.encode(), "text/plain"))
                 for i in range(11)]
        r = client.post("/api/v1/projects/upload/sources", files=files)
        assert r.status_code == 422

    def test_methodichka_upload(self, client):
        text = ("Методические указания\n"
                "Объём курсовой работы составляет 30 страниц.\n"
                "Работа должна содержать 3 главы.\n" + "требования " * 60)
        r = client.post(
            "/api/v1/projects/upload/methodichka",
            files={"file": ("m.txt", text.encode(), "text/plain")},
        )
        assert r.status_code == 200
        assert r.json()["found"]

    def test_methodichka_bad_format(self, client):
        r = client.post(
            "/api/v1/projects/upload/methodichka",
            files={"file": ("m.exe", b"x" * 900, "application/octet-stream")},
        )
        assert r.status_code == 422


class TestTitlePageTrap:
    """Титульный лист не должен становиться заголовком.

    Реальные файлы пользователя вскрыли то, чего не было в синтетике:
    у магистерской диссертации заголовком стало «МОСКОВСКИЙ
    ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ», а автором — заглушка бланка
    «И.О. Фамилия».
    """

    DIPLOMA = "\n".join([
        "МОСКОВСКИЙ ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ",
        "МФЮА",
        "Кафедра гражданско-правовых дисциплин",
        "Направление / специальность 40.04.01 Юриспруденция",
        "К ЗАЩИТЕ",
        "(РЕКОМЕНДОВАНО / НЕ РЕКОМЕНДОВАНО)",
        "Заведующий кафедрой",
        "(подпись) (И.О. Фамилия)",
        "«_____» ____________ 2026 г.",
        "МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ",
        "на тему:",
        "«СУБСИДИАРНАЯ ОТВЕТСТВЕННОСТЬ КОНТРОЛИРУЮЩИХ ЛИЦ»",
        "Обучающийся:",
        "(Ф.И.О.)",
        "Москва – 2026",
    ]) + "\n" + "содержательный текст работы. " * 40

    DISSERTATION = "\n".join([
        "ИВАНОВСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ",
        "На правах рукописи",
        "КОКОЛОВ СЕРГЕЙ БОРИСОВИЧ",
        "",
        "УПРАВЛЕНИЕ ЭФФЕКТИВНОСТЬЮ РЕГИОНАЛЬНОГО",
        "НАЛОГООБЛОЖЕНИЯ В РЫНОЧНЫХ УСЛОВИЯХ",
        "(НА ПРИМЕРЕ ИВАНОВСКОЙ ОБЛАСТИ)",
        "",
        "Научный руководитель",
        "доктор технических наук, профессор",
        "К.М. Пирогов",
        "Иваново, 2002",
    ]) + "\n" + "текст диссертации. " * 40

    def test_topic_marker_wins_over_university(self):
        assert guess_title(self.DIPLOMA, "d.pdf") == \
            "СУБСИДИАРНАЯ ОТВЕТСТВЕННОСТЬ КОНТРОЛИРУЮЩИХ ЛИЦ"

    def test_placeholder_author_ignored(self):
        assert "И.О. Фамилия" not in guess_authors(self.DIPLOMA)

    def test_multiline_caps_title_joined(self):
        title = guess_title(self.DISSERTATION, "f.pdf")
        assert title.startswith("УПРАВЛЕНИЕ ЭФФЕКТИВНОСТЬЮ РЕГИОНАЛЬНОГО")
        assert "ИВАНОВСКОЙ ОБЛАСТИ" in title

    def test_fullname_caps_not_a_title(self):
        assert "КОКОЛОВ" not in guess_title(self.DISSERTATION, "f.pdf")

    def test_blank_form_lines_skipped(self):
        assert "подпись" not in guess_title(self.DIPLOMA, "d.pdf").lower()
        assert "_____" not in guess_title(self.DIPLOMA, "d.pdf")

    def test_quoted_midtext_not_taken_as_title(self):
        text = ("Бизнес-модели для IT-стартапов\n"
                "«Одноклассники» сегодня это не только социальный ресурс, "
                "но и известный бренд, который развивается\n"
                + "текст статьи. " * 40)
        assert guess_title(text, "f.pdf") == "Бизнес-модели для IT-стартапов"

    def test_title_with_colon_not_merged_with_body(self):
        text = ("Как открыть кофейню с нуля: от идеи до воплощения\n"
                "Открытие кофейни требует продуманного подхода к выбору "
                "помещения и оборудования\n" + "далее текст. " * 40)
        assert guess_title(text, "f.pdf") == \
            "Как открыть кофейню с нуля: от идеи до воплощения"


class TestMeaningfulExcerpt:
    """Титульник и оглавление не должны попадать в промпт.

    Загруженная диссертация на 187 тысяч знаков дала ноль тезисов:
    в промпт уходили первые 3500 знаков, то есть название вуза,
    подписи и список глав с номерами страниц.
    """

    DOC = (
        "МОСКОВСКИЙ УНИВЕРСИТЕТ\n"
        "МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ\n"
        "Москва – 2026\n"
        "СОДЕРЖАНИЕ\n"
        "ВВЕДЕНИЕ ........... 3\n"
        "ГЛАВА 1. ПОНЯТИЕ ОТВЕТСТВЕННОСТИ ........... 8\n"
        "1.1. Правовая природа ........... 8\n"
        "ЗАКЛЮЧЕНИЕ ........... 70\n"
        "ВВЕДЕНИЕ\n"
        "Субсидиарная ответственность контролирующих должника лиц "
        "выступает ключевым инструментом защиты интересов кредиторов. "
        + "Существенное влияние оказало Постановление Пленума. " * 60
    )

    def test_starts_from_introduction(self):
        from app.modules.sources.user_upload import meaningful_excerpt
        out = meaningful_excerpt(self.DOC, 1000)
        assert out.startswith("Субсидиарная ответственность")

    def test_drops_toc_lines(self):
        from app.modules.sources.user_upload import meaningful_excerpt
        out = meaningful_excerpt(self.DOC, 1000)
        assert "..........." not in out
        assert "МОСКОВСКИЙ УНИВЕРСИТЕТ" not in out

    def test_short_text_returned_whole(self):
        from app.modules.sources.user_upload import meaningful_excerpt
        assert meaningful_excerpt("короткий текст", 1000) == "короткий текст"

    def test_no_markers_falls_back_to_head(self):
        from app.modules.sources.user_upload import meaningful_excerpt
        text = "просто текст без разделов. " * 300
        assert meaningful_excerpt(text, 500).startswith("просто текст")

    def test_prompt_uses_excerpt_for_uploads(self):
        from app.modules.sources.grounding import format_sources_for_prompt
        from app.modules.sources.user_upload import build_source
        s = build_source("d.pdf", self.DOC, title="Тема").source
        out = format_sources_for_prompt([s])
        assert "Субсидиарная ответственность контролирующих" in out
        assert "МОСКОВСКИЙ УНИВЕРСИТЕТ" not in out
