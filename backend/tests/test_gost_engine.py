"""Тесты ГОСТ-движка.

Цель — зафиксировать правила оформления, выверенные визуально в оригинале
на JS. Если кто-то «поправит» отступ, размер шрифта или разрыв страницы,
тест должен упасть.
"""

from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.modules.documents.gost_engine import (
    FONT,
    fix_dashes,
    generate_fragment_docx,
    generate_full_docx,
    parse_markdown_table,
    split_paragraphs,
    strip_duplicate_heading,
)


def load(data: bytes) -> Document:
    return Document(BytesIO(data))


def texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


# ------------------------------------------------------------ fix_dashes

class TestFixDashes:
    def test_number_range_gets_plain_hyphen(self):
        # Диапазон — это не пауза в речи, пробелов быть не должно.
        assert fix_dashes("2025 — 2026 гг.") == "2025-2026 гг."
        assert fix_dashes("ст. 61.11 – 61.12") == "ст. 61.11-61.12"

    def test_em_dash_becomes_en_dash(self):
        assert fix_dashes("эффект — преодоление") == "эффект – преодоление"

    def test_range_processed_before_general_replace(self):
        # Оба случая в одной строке: диапазон не должен превратиться в тире.
        assert fix_dashes("в 2020—2021 гг. вывод — очевиден") == (
            "в 2020-2021 гг. вывод – очевиден"
        )

    def test_empty_input(self):
        assert fix_dashes("") == ""
        assert fix_dashes(None) == ""


# ------------------------------------------------------- split_paragraphs

class TestSplitParagraphs:
    def test_splits_on_blank_lines(self):
        assert split_paragraphs("Первый.\n\nВторой.") == ["Первый.", "Второй."]

    def test_ignores_extra_whitespace_and_empties(self):
        assert split_paragraphs("A.\n   \n\nB.\n\n\n") == ["A.", "B."]

    def test_empty(self):
        assert split_paragraphs("") == []
        assert split_paragraphs(None) == []


# -------------------------------------------------- strip_duplicate_heading

class TestStripDuplicateHeading:
    def test_removes_exact_duplicate(self):
        assert strip_duplicate_heading("Введение\nТекст.", "Введение") == "Текст."

    def test_removes_when_heading_is_prefix(self):
        # Заголовок "1.1" против первой строки "1.1. Название раздела".
        out = strip_duplicate_heading("1.1. Понятие\nТекст.", "1.1")
        assert out == "Текст."

    def test_removes_bare_vvedenie_regardless_of_heading(self):
        assert strip_duplicate_heading("ВВЕДЕНИЕ\nТекст.", "Что-то") == "Текст."

    def test_keeps_long_first_line(self):
        # Длинная первая строка — это уже текст, отрезать нельзя.
        long_line = "Э" * 200
        raw = f"{long_line}\nВторая."
        assert strip_duplicate_heading(raw, "Введение") == raw

    def test_keeps_text_when_no_duplicate(self):
        raw = "Совсем другое начало.\nДальше."
        assert strip_duplicate_heading(raw, "Введение") == raw


# ------------------------------------------------------ parse_markdown_table

class TestParseMarkdownTable:
    def test_basic_table(self):
        md = "|Критерий|Описание|\n|---|---|\n|Субъект|КДЛ|"
        assert parse_markdown_table(md) == [
            ["Критерий", "Описание"],
            ["Субъект", "КДЛ"],
        ]

    def test_merges_wrapped_cell_continuation(self):
        """Ключевой случай: модель перенесла текст ячейки без '|'.

        Наивный парсер потерял бы продолжение и сдвинул данные —
        это тихая порча данных, а не косметика.
        """
        md = (
            "|A|B|\n"
            "|---|---|\n"
            "|x|начало,\n"
            "продолжение|\n"
            "|y|z|"
        )
        rows = parse_markdown_table(md)
        assert rows[1] == ["x", "начало, продолжение"]
        assert rows[2] == ["y", "z"]

    @pytest.mark.xfail(
        reason=(
            "ИЗВЕСТНЫЙ ДЕФЕКТ, унаследован от оригинала на JS (проверено "
            "сравнением: JS даёт ровно тот же результат). Зона таблицы "
            "обрезается по ПОСЛЕДНЕЙ строке, начинающейся с '|', поэтому "
            "перенос в САМОЙ ПОСЛЕДНЕЙ ячейке отбрасывается и текст молча "
            "теряется. Внутри таблицы перенос склеивается верно. "
            "Чинить вместе с заказчиком: см. docs/GOST_PORT.md"
        ),
        strict=True,
    )
    def test_trailing_wrapped_cell_is_lost(self):
        md = (
            "|Критерий|Описание|\n"
            "|---|---|\n"
            "|Субъект|Контролирующее лицо,\n"
            "которое давало указания|"
        )
        rows = parse_markdown_table(md)
        assert rows[1][1] == "Контролирующее лицо, которое давало указания"

    def test_extra_columns_are_joined_not_dropped(self):
        md = "|A|B|\n|---|---|\n|1|2|3|"
        rows = parse_markdown_table(md)
        assert rows[1] == ["1", "2 | 3"]

    def test_separator_row_removed(self):
        md = "|A|B|\n|:---|---:|\n|1|2|"
        rows = parse_markdown_table(md)
        assert len(rows) == 2

    def test_returns_none_for_garbage(self):
        assert parse_markdown_table(None) is None
        assert parse_markdown_table("обычный текст") is None
        assert parse_markdown_table("|только одна строка|") is None


# ------------------------------------------------------------ вёрстка

class TestFragmentDocx:
    def test_body_formatting_matches_gost(self):
        data = generate_fragment_docx(
            title="1.1. Понятие", text="Первый абзац.\n\nВторой абзац."
        )
        doc = load(data)
        body = [p for p in doc.paragraphs if p.text == "Первый абзац."][0]

        assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert body.paragraph_format.first_line_indent == Pt(36)  # 720 twips
        assert body.paragraph_format.line_spacing == 1.5
        assert body.paragraph_format.space_before == Pt(0)
        assert body.paragraph_format.space_after == Pt(0)
        run = body.runs[0]
        assert run.font.name == FONT
        assert run.font.size == Pt(14)
        assert not run.font.bold

    def test_h1_is_uppercase_centered_not_bold(self):
        data = generate_fragment_docx(title="ВВЕДЕНИЕ", text="Текст.")
        doc = load(data)
        head = [p for p in doc.paragraphs if "ВВЕДЕНИЕ" in p.text][0]
        assert head.alignment == WD_ALIGN_PARAGRAPH.CENTER
        run = [r for r in head.runs if r.text][0]
        assert run.text == "ВВЕДЕНИЕ"
        assert not run.font.bold, "H1 по требованиям НЕ жирный"
        assert run.font.size == Pt(16)

    def test_h1_detected_by_title_name(self):
        # "ЗАКЛЮЧЕНИЕ" в списке H1_TITLES → капс и центр без явного флага.
        doc = load(generate_fragment_docx(title="Заключение", text="Т."))
        head = [p for p in doc.paragraphs if p.text.strip() == "ЗАКЛЮЧЕНИЕ"]
        assert head and head[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_h2_is_not_uppercased_nor_centered(self):
        doc = load(generate_fragment_docx(title="1.1. Понятие", text="Т."))
        head = [p for p in doc.paragraphs if p.text == "1.1. Понятие"][0]
        assert head.alignment != WD_ALIGN_PARAGRAPH.CENTER
        assert head.runs[0].font.size == Pt(14)

    def test_page_margins_gost(self):
        """Поля сверяем в twips: Word хранит их именно так.

        Прямое сравнение с Cm() даёт ложное расхождение из-за округления
        EMU → twips (Cm(3) = 1080000 EMU, но 1701 twips = 1080135 EMU).
        """
        doc = load(generate_fragment_docx(title="Т", text="Т."))
        s = doc.sections[0]
        assert s.left_margin.twips == Cm(3).twips, "слева 3 см — переплёт"
        assert s.right_margin.twips == Cm(1.5).twips
        assert s.top_margin.twips == Cm(2).twips
        assert s.bottom_margin.twips == Cm(2).twips

    def test_duplicate_heading_not_rendered_twice(self):
        doc = load(
            generate_fragment_docx(title="Введение", text="Введение\nНастоящий текст.")
        )
        assert texts(doc).count("ВВЕДЕНИЕ") == 1
        assert "Настоящий текст." in texts(doc)

    def test_table_with_caption(self):
        data = generate_fragment_docx(
            title="1.1. Понятие",
            text="Текст.",
            table_markdown="|A|B|\n|---|---|\n|1|2|",
            table_number="1.1",
            table_title="Критерии",
        )
        doc = load(data)
        assert any("Таблица № 1.1. Критерии" in t for t in texts(doc))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "A"
        assert table.rows[1].cells[1].text == "2"

    def test_table_font_is_11pt(self):
        data = generate_fragment_docx(
            title="Т", text="Т.", table_markdown="|A|B|\n|---|---|\n|1|2|"
        )
        doc = load(data)
        run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
        assert run.font.size == Pt(11)

    def test_table_header_centered_body_left(self):
        data = generate_fragment_docx(
            title="Т", text="Т.", table_markdown="|A|B|\n|---|---|\n|1|2|"
        )
        t = load(data).tables[0]
        assert t.rows[0].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert t.rows[1].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_reference_sentence_appended(self):
        data = generate_fragment_docx(
            title="1.1",
            text="Текст.",
            table_markdown="|A|B|\n|---|---|\n|1|2|",
            table_number="1.1",
            table_title="Крит",
            reference_sentence="Данные приведены в таблице 1.1.",
        )
        assert "Данные приведены в таблице 1.1." in texts(load(data))

    def test_dashes_normalized_in_output(self):
        doc = load(generate_fragment_docx(title="Т", text="За 2020 — 2021 гг."))
        assert any("2020-2021" in t for t in texts(doc))


class TestFullDocx:
    @pytest.fixture
    def full(self) -> Document:
        return load(
            generate_full_docx(
                topic="Субсидиарная ответственность",
                introduction="Текст введения.",
                sections=[
                    {"number": "1.1", "text": "Текст 1.1."},
                    {"number": "1.2", "text": "Текст 1.2."},
                    {
                        "number": "2.1",
                        "text": "Текст 2.1.",
                        "table": {
                            "number": "2.1",
                            "title": "Крит",
                            "markdown": "|A|B|\n|---|---|\n|1|2|",
                        },
                    },
                ],
                conclusion="Текст заключения.",
                chapter_titles={"1": "Теория", "2": "Практика"},
                section_titles={"1.1": "Понятие", "1.2": "Признаки", "2.1": "Анализ"},
            )
        )

    def test_mandatory_sections_present(self, full):
        t = texts(full)
        for expected in ("СОДЕРЖАНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
            assert expected in t, f"нет раздела {expected}"

    def test_chapters_grouped_by_number_prefix(self, full):
        t = texts(full)
        # 1.1 и 1.2 → одна ГЛАВА 1; 2.1 → ГЛАВА 2.
        assert t.count("ГЛАВА 1. ТЕОРИЯ") == 1
        assert t.count("ГЛАВА 2. ПРАКТИКА") == 1

    def test_contents_lists_chapters_and_sections(self, full):
        t = texts(full)
        idx = t.index("СОДЕРЖАНИЕ")
        contents = t[idx: idx + 12]
        assert "ВВЕДЕНИЕ" in contents
        assert any("1.1. Понятие" in line for line in contents)
        assert any("ГЛАВА 1. Теория" in line for line in contents)

    def test_contents_sections_are_indented(self, full):
        para = [p for p in full.paragraphs if p.text == "1.1. Понятие"][0]
        assert para.paragraph_format.left_indent is not None

    def test_section_order_preserved(self, full):
        """Порядок тела работы.

        ВВЕДЕНИЕ и ЗАКЛЮЧЕНИЕ встречаются дважды: строкой в СОДЕРЖАНИИ и
        собственно заголовком. Поэтому берём ПОСЛЕДНЕЕ вхождение —
        первое относится к оглавлению.
        """
        t = texts(full)
        intro_heading = len(t) - 1 - t[::-1].index("ВВЕДЕНИЕ")
        concl_heading = len(t) - 1 - t[::-1].index("ЗАКЛЮЧЕНИЕ")
        assert intro_heading < t.index("Текст 1.1.")
        assert t.index("Текст 1.1.") < t.index("Текст 2.1.")
        assert t.index("Текст 2.1.") < concl_heading

    def test_table_rendered_in_full_doc(self, full):
        assert len(full.tables) == 1
        assert any("Таблица № 2.1. Крит" in t for t in texts(full))

    def test_empty_sections_do_not_crash(self):
        data = generate_full_docx(
            introduction="Вступление.", sections=[], conclusion="Итог."
        )
        t = texts(load(data))
        assert "ВВЕДЕНИЕ" in t and "ЗАКЛЮЧЕНИЕ" in t
