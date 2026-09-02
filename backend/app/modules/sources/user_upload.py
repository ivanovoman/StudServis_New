"""Источники, которые пользователь загрузил сам.

Автопоиск покрывает не всё. По узкой теме OpenAlex и КиберЛенинка
возвращают статьи, совпавшие с темой лишь общими словами, — проверка
достаточности покрытия честно об этом сообщает и советует подобрать
литературу вручную. Этот модуль и есть ручной путь.

Что важнее: у научного руководителя часто есть список обязательной
литературы, а у студента — методичка с требованием опереться на
конкретные работы. Такие источники нельзя ни найти автоматически, ни
проигнорировать.

## Приоритет

Загруженные пользователем источники не проходят отбор по релевантности
и не могут быть вытеснены найденными автоматически. Пользователь уже
решил, что эти работы нужны; наше дело — разобрать их, а не спорить.

## Про полный текст

В отличие от автопоиска, здесь почти всегда есть полный текст — файл
загружен целиком. Это лучший материал для разбора: именно на полных
текстах КиберЛенинки анализ начал видеть научную полемику, которой в
абстрактах не видно.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from app.modules.sources.openalex import Source

#: Что принимаем от пользователя.
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".rtf"}

#: Больше — почти наверняка не статья, а сканированная книга целиком.
MAX_FILE_BYTES = 25 * 1024 * 1024

#: Короче — файл пустой или это скан без текстового слоя.
MIN_USABLE_CHARS = 400

#: Текст обрезается при укладке в промпт, но хранится целиком:
#: он ещё пойдёт в RAG, где длина полезна.
MAX_STORED_CHARS = 200_000


class UploadError(ValueError):
    """Файл не удалось использовать. Сообщение предназначено пользователю."""


@dataclass
class UploadedSource:
    """Результат разбора одного загруженного файла."""

    source: Source
    filename: str
    chars: int
    warnings: list[str]

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "chars": self.chars,
            "warnings": self.warnings,
            "source": self.source.to_dict(),
        }


# ---------------------------------------------------------------- извлечение

def extract_txt(data: bytes) -> str:
    """Прочитать текстовый файл, угадав кодировку.

    Русские тексты приходят и в UTF-8, и в CP1251 — второе особенно
    часто у файлов, сохранённых из старого Word. Отличить их надёжно
    можно по тому, даёт ли декодирование осмысленную кириллицу.
    """
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            # BOM снимаем отдельно: некоторые редакторы пишут его дважды,
            # и utf-8-sig убирает только первый.
            return data.decode(encoding).lstrip("\ufeff")
        except UnicodeDecodeError:
            continue

    best_text = ""
    best_score = -1.0
    for encoding in ("cp1251", "koi8-r", "cp866", "iso8859-5"):
        try:
            text = data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
        letters = sum(1 for c in text if c.isalpha())
        if not letters:
            continue
        cyrillic = sum(1 for c in text if "а" <= c.lower() <= "я")
        score = cyrillic / letters
        if score > best_score:
            best_score, best_text = score, text

    if best_text:
        return best_text
    return data.decode("utf-8", "replace")


def extract_pdf_bytes(data: bytes) -> str:
    """Достать текст из PDF.

    Переносы по слогам склеиваются: PDF рвёт слова дефисом в конце
    строки, и без склейки поиск по словам ломается.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise UploadError("на сервере нет библиотеки для чтения PDF") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise UploadError(
            "не удалось прочитать PDF — файл повреждён или защищён паролем"
        ) from exc

    text = "\n\n".join(pages)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text


def extract_docx_bytes(data: bytes) -> str:
    """Достать текст из DOCX, включая таблицы."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise UploadError("на сервере нет библиотеки для чтения DOCX") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise UploadError(
            "не удалось прочитать DOCX — возможно, это старый формат .doc, "
            "пересохраните файл как .docx"
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_rtf(data: bytes) -> str:
    """Грубо развернуть RTF.

    Полноценный парсер RTF нам не нужен: достаточно снять управляющие
    последовательности и раскодировать кириллицу из \\'xx.
    """
    text = data.decode("cp1251", "replace")
    text = re.sub(r"\\'([0-9a-fA-F]{2})",
                  lambda m: bytes([int(m.group(1), 16)]).decode(
                      "cp1251", "replace"), text)
    text = re.sub(r"\\par[d]?\b", "\n", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"[ \t]{2,}", " ", text)


def extract_text(filename: str, data: bytes) -> str:
    """Извлечь текст из файла по его расширению."""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise UploadError(
            f"формат {suffix or 'без расширения'} не поддерживается; "
            f"допустимы: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    if not data:
        raise UploadError("файл пустой")
    if len(data) > MAX_FILE_BYTES:
        raise UploadError(
            f"файл больше {MAX_FILE_BYTES // (1024 * 1024)} МБ — "
            "загрузите отдельную статью, а не сборник"
        )

    if suffix == ".pdf":
        return extract_pdf_bytes(data)
    if suffix == ".docx":
        return extract_docx_bytes(data)
    if suffix == ".rtf":
        return extract_rtf(data)
    return extract_txt(data)


# ------------------------------------------------------------- метаданные

RE_YEAR = re.compile(r"\b(19[5-9]\d|20[0-4]\d)\b")

RE_DOI = re.compile(r"\b10\.\d{4,9}/[^\s\"'<>,;]+", re.I)

RE_UDK = re.compile(r"^\s*(?:УДК|ББК|ГРНТИ)[\s\d.:()\-–/]+$", re.I)

#: «Иванов И. И.» или «И. И. Иванов» — типовая подпись автора статьи.
RE_AUTHOR = re.compile(
    r"([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s?[А-ЯЁ]\.)"
    r"|([А-ЯЁ]\.\s?[А-ЯЁ]\.\s+[А-ЯЁ][а-яё]+)"
)

#: Строки, которые заголовком быть не могут.
RE_SERVICE_LINE = re.compile(
    r"^\s*(?:УДК|ББК|ГРНТИ|DOI|аннотация|ключевые\s+слова|abstract|"
    r"keywords|страниц|стр\.|©|удк)\b", re.I)

#: Титульный лист. У курсовых и дипломов первые строки — это название
#: вуза, кафедра и гриф «допустить к защите», а настоящая тема идёт
#: ниже. Без этого фильтра заголовком становился «МОСКОВСКИЙ
#: ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ».
RE_TITLE_PAGE = re.compile(
    r"(?:университет|институт|академия|факультет|кафедр|колледж|"
    r"министерств|образовани|автономн\w+\s+учрежден|"
    r"выпускная\s+квалификационная|курсовая\s+работа|"
    r"дипломная\s+работа|магистерская\s+диссертация|"
    r"допустить\s+к\s+защите|направление\s+подготовки|"
    r"выполнил|руководител|рецензент|москва\s*,?\s*20\d\d)", re.I)

#: Маркер темы на титульном листе: «на тему:», «Тема:». Следующая
#: строка — настоящее название работы, а не название вуза.
RE_TOPIC_MARKER = re.compile(
    r"^\s*(?:на\s+тему|тема(?:\s+работы)?)\s*:?\s*$", re.I)

#: Строка-бланк: подчёркивания для вписывания от руки, «(подпись)»,
#: «(Ф.И.О.)», даты с пропусками.
RE_BLANK_LINE = re.compile(
    r"_{3,}|^\s*\((?:подпись|Ф\.?И\.?О\.?|И\.?О\.?\s*Фамилия|"
    r"ученая\s+степень[^)]*|РЕКОМЕНДОВАНО[^)]*)\)\s*$", re.I)

#: ФИО целиком капсом: «КОКОЛОВ СЕРГЕЙ БОРИСОВИЧ». На титуле
#: диссертации это автор, а не заголовок.
RE_FULLNAME_CAPS = re.compile(
    r"^[А-ЯЁ]{3,}\s+[А-ЯЁ]{3,}\s+[А-ЯЁ]{4,}(?:ИЧ|НА|ВНА|ОВИЧ|ЕВИЧ)$")

#: Шаблонные заглушки из бланков, которые нельзя считать автором.
RE_PLACEHOLDER_AUTHOR = re.compile(
    r"^(?:И\.\s?О\.\s+Фамилия|Фамилия\s+И\.\s?О\.|"
    r"[А-ЯЁ]\.\s?[А-ЯЁ]\.\s+Фамилия)$", re.I)


def _clean_line(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip(" \t\u00a0*#—–-")


def _is_title_candidate(line: str) -> bool:
    """Может ли строка быть частью заголовка работы."""
    if not line or len(line) < 6 or len(line) > 250:
        return False
    if RE_SERVICE_LINE.match(line) or RE_UDK.match(line):
        return False
    if RE_BLANK_LINE.search(line) or RE_DOI.search(line):
        return False
    if RE_TITLE_PAGE.search(line) or RE_TOPIC_MARKER.match(line):
        return False
    if RE_AUTHOR.fullmatch(line.strip()) or RE_FULLNAME_CAPS.match(line):
        return False
    letters = sum(1 for c in line if c.isalpha())
    return letters >= len(line) * 0.5


def _caps_ratio(line: str) -> float:
    letters = [c for c in line if c.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for c in letters if c.isupper()) / len(letters)


def _gather_block(lines: list[str], start: int, max_lines: int = 4) -> str:
    """Собрать заголовок, разбитый на несколько строк.

    В PDF длинный заголовок переносится: «УПРАВЛЕНИЕ ЭФФЕКТИВНОСТЬЮ
    РЕГИОНАЛЬНОГО / НАЛОГООБЛОЖЕНИЯ В РЫНОЧНЫХ УСЛОВИЯХ / (НА ПРИМЕРЕ
    ИВАНОВСКОЙ ОБЛАСТИ)». Продолжением считаем только соседние строки
    того же регистра, идущие без пустой строки между ними.
    """
    first = lines[start]
    parts = [first]
    caps = _caps_ratio(first) > 0.7

    # Заголовок, кончающийся точкой, продолжения не имеет.
    if first.endswith((".", "!", "?")):
        return first.strip("«»\"' ")

    for line in lines[start + 1:start + max_lines]:
        if not line:
            break
        if not _is_title_candidate(line):
            break
        # Регистр должен совпадать: капс-заголовок не продолжается прозой.
        if (_caps_ratio(line) > 0.7) != caps:
            break
        # Продолжением заголовка считаем только строку, начатую со
        # строчной буквы. Заглавная (в том числе после открывающей
        # кавычки) — это уже новое предложение, начало текста статьи.
        if not caps:
            head = line.lstrip("«\"'(")[:1]
            if head.isupper() and not first.endswith(","):
                break
        parts.append(line)
        if line.endswith((".", "»", '"')):
            break

    title = " ".join(parts).strip()
    return title.strip("«»\"' ")


def guess_title(text: str, fallback: str) -> str:
    """Угадать заголовок работы по первым строкам.

    Порядок проверок отражает надёжность признака:

    1. Явный маркер «на тему:» — на титульном листе курсовой или
       диссертации следующая строка и есть тема.
    2. Строка в кавычках среди шапки.
    3. Блок строк капслоком — типовое оформление заголовка.
    4. Первая содержательная строка.
    5. Имя файла.

    Без первых трёх заголовком становилось название вуза: реальный
    файл диссертации давал «МОСКОВСКИЙ ФИНАНСОВО-ЮРИДИЧЕСКИЙ
    УНИВЕРСИТЕТ» вместо темы работы.
    """
    lines = [_clean_line(l) for l in text.split("\n")[:80]]

    # 1. Явный маркер темы.
    for i, line in enumerate(lines):
        if not RE_TOPIC_MARKER.match(line):
            continue
        for j in range(i + 1, min(i + 4, len(lines))):
            if _is_title_candidate(lines[j]):
                return _gather_block(lines, j)

    # 2. Название целиком в кавычках. Требуем и закрывающую кавычку:
    # иначе правило ловит цитату, которой начинается абзац текста.
    for i, line in enumerate(lines[:40]):
        stripped = line.strip()
        if (len(stripped) > 20
                and stripped.startswith(("«", '"'))
                and stripped.endswith(("»", '"'))
                and _is_title_candidate(stripped)):
            return stripped.strip("«»\"' ")

    # 3. Блок капслоком.
    for i, line in enumerate(lines[:50]):
        if _is_title_candidate(line) and _caps_ratio(line) > 0.7:
            return _gather_block(lines, i)

    # 4. Первая содержательная строка.
    for i, line in enumerate(lines):
        if _is_title_candidate(line) and len(line) >= 12:
            return _gather_block(lines, i)

    return Path(fallback).stem.replace("_", " ").strip() or fallback


def guess_year(text: str) -> int | None:
    """Найти год издания.

    Ищем в начале и в конце: в шапке статьи и в выходных данных. Из
    нескольких кандидатов берём наибольший — год издания не может быть
    меньше года цитируемых работ.
    """
    head = text[:3000]
    tail = text[-3000:]
    years = [int(y) for y in RE_YEAR.findall(head + "\n" + tail)]
    return max(years) if years else None


def guess_authors(text: str, limit: int = 3) -> list[str]:
    """Собрать фамилии авторов из шапки статьи."""
    head = text[:2500]
    found: list[str] = []
    for match in RE_AUTHOR.finditer(head):
        name = _clean_line(match.group(0))
        # «И.О. Фамилия» — заглушка из бланка титульного листа.
        if RE_PLACEHOLDER_AUTHOR.match(name):
            continue
        if name and name not in found:
            found.append(name)
        if len(found) >= limit:
            break
    return found


def guess_doi(text: str) -> str | None:
    match = RE_DOI.search(text[:4000])
    if not match:
        return None
    return match.group(0).rstrip(".,;")


# ------------------------------------------------------------------ разбор

def build_source(
    filename: str,
    text: str,
    *,
    title: str | None = None,
    year: int | None = None,
    authors: list[str] | None = None,
) -> UploadedSource:
    """Собрать Source из извлечённого текста.

    Явно переданные пользователем поля всегда важнее угаданных: он
    видел документ своими глазами, а мы — только регулярки.
    """
    warnings: list[str] = []
    text = text.strip()

    if len(text) < MIN_USABLE_CHARS:
        raise UploadError(
            f"из файла удалось извлечь всего {len(text)} знаков. "
            "Похоже, это скан без текстового слоя — распознайте его "
            "или загрузите текстовую версию"
        )

    if len(text) > MAX_STORED_CHARS:
        text = text[:MAX_STORED_CHARS]
        warnings.append(
            f"текст обрезан до {MAX_STORED_CHARS} знаков")

    resolved_title = (title or "").strip() or guess_title(text, filename)
    resolved_year = year or guess_year(text)
    resolved_authors = authors or guess_authors(text)

    if not title:
        warnings.append(f"заголовок определён автоматически: «{resolved_title}»")
    if resolved_year is None:
        warnings.append("год издания не найден — укажите его вручную")
    if not resolved_authors:
        warnings.append("авторы не распознаны")

    source = Source(
        title=resolved_title,
        year=resolved_year,
        doi=guess_doi(text),
        abstract=text[:1500],
        url="",
        authors=resolved_authors,
        venue=None,
        cited_by=0,
        is_oa=True,
        provider="user_upload",
        # Пользователь выбрал этот источник сам — он релевантен по
        # определению, и отбор по релевантности его не касается.
        relevance=1.0,
    )
    source.fulltext = text

    return UploadedSource(
        source=source,
        filename=filename,
        chars=len(text),
        warnings=warnings,
    )


def load_uploaded_source(
    filename: str,
    data: bytes,
    *,
    title: str | None = None,
    year: int | None = None,
    authors: list[str] | None = None,
) -> UploadedSource:
    """Полный разбор загруженного файла: текст, метаданные, Source."""
    text = extract_text(filename, data)
    return build_source(filename, text, title=title, year=year,
                        authors=authors)


def merge_with_found(
    uploaded: Iterable[Source],
    found: Iterable[Source],
    *,
    limit: int = 6,
) -> list[Source]:
    """Соединить загруженные источники с найденными автоматически.

    Загруженные идут первыми и не вытесняются: пользователь уже решил,
    что они нужны. Автопоиск лишь дополняет список до нужного объёма.
    Если пользователь загрузил больше лимита — берём все его, потому
    что молча выбрасывать выбранное пользователем нельзя.
    """
    from app.modules.sources.openalex import normalize_title

    result = list(uploaded)
    seen = {normalize_title(s.title) for s in result if s.title}

    for source in found:
        if len(result) >= limit:
            break
        key = normalize_title(source.title)
        if key and key in seen:
            continue
        seen.add(key)
        result.append(source)

    return result


# --------------------------------------------------------- полезный фрагмент

RE_INTRO_START = re.compile(
    r"^\s*(?:ВВЕДЕНИЕ|Введение|ANNOTATION|Аннотация|"
    r"АННОТАЦИЯ|Abstract)\s*$", re.M)

RE_TOC_START = re.compile(
    r"^\s*(?:СОДЕРЖАНИЕ|ОГЛАВЛЕНИЕ|Содержание|Оглавление)\s*$", re.M)

#: Строка оглавления: «1.2. Название раздела ......... 15».
RE_TOC_LINE = re.compile(r"\.{3,}\s*\d+\s*$|\s{3,}\d{1,3}\s*$", re.M)


def meaningful_excerpt(text: str, limit: int = 3500) -> str:
    """Вырезать содержательный фрагмент, пропустив титульник и оглавление.

    Загруженная курсовая или диссертация начинается с титульного листа
    и оглавления. Если слепо взять начало файла, в промпт уйдут
    название вуза, подписи и список глав с номерами страниц — модель
    не найдёт там ни одной мысли.

    Реальный случай: диссертация на 187 тысяч знаков дала ноль тезисов
    именно поэтому.
    """
    start = 0

    # Введение — самое полезное начало: там постановка проблемы.
    intro = RE_INTRO_START.search(text[:60_000])
    if intro:
        start = intro.end()
    else:
        toc = RE_TOC_START.search(text[:20_000])
        if toc:
            # После оглавления: ищем конец списка с номерами страниц.
            tail = text[toc.end():toc.end() + 30_000]
            last = None
            for match in RE_TOC_LINE.finditer(tail):
                last = match
            start = toc.end() + (last.end() if last else 0)

    # Короткий документ без служебной шапки берём целиком.
    if start == 0 and len(text) <= limit:
        return text

    excerpt = text[start:start + limit * 2].strip()

    # Строки оглавления могли уцелеть — выбрасываем их.
    lines = [l for l in excerpt.split("\n") if not RE_TOC_LINE.search(l)]
    excerpt = "\n".join(lines).strip()

    return excerpt[:limit] if excerpt else text[:limit]
