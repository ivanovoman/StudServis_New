"""ГОСТ-вёрстка DOCX.

Порт api/docxExport.js на python-docx. Правила оформления сохранены
один в один — они были выверены визуально через рендеринг в PDF, поэтому
здесь ничего не «улучшено по вкусу».

Требования, зафиксированные оригиналом:
- H1 (СОДЕРЖАНИЕ, ВВЕДЕНИЕ, ГЛАВА N, ЗАКЛЮЧЕНИЕ, СПИСОК ЛИТЕРАТУРЫ) —
  КАПС, БЕЗ жирности, по центру, разрыв страницы перед заголовком,
  интервалы before/after = 0.
- H2 (разделы 1.1, 1.2) — обычный шрифт, без разрыва страницы, before/after = 0.
- Основной текст — Times New Roman 14pt, межстрочный 1.5, выравнивание по
  ширине, отступ первой строки 1.25 см (720 DXA), before/after = 0.
- Таблицы — без заливки, 11pt, шапка по центру, тело по левому краю,
  границы 1.5pt.
- Поля: верх 2 см, низ 2 см, слева 3 см (переплёт), справа 1.5 см. A4.
- Отступы между блоками делаются ПУСТЫМИ АБЗАЦАМИ, а не spacing —
  так у заголовков и текста before/after остаются нулевыми.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Iterable, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

# ---- Единицы и константы оформления ----
# python-docx оперирует Pt/Cm/Twips напрямую, поэтому DXA-константы
# оригинала переведены в осмысленные единицы.
FONT = "Times New Roman"
BODY_SIZE = Pt(14)
TABLE_SIZE = Pt(11)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)

LINE_SPACING = 1.5
FIRST_LINE_INDENT = Twips(720)   # 1.25 см — красная строка
CONTENTS_INDENT = Twips(720)     # отступ разделов в СОДЕРЖАНИИ

PAGE_MARGIN_TOP = Cm(2)
PAGE_MARGIN_BOTTOM = Cm(2)
PAGE_MARGIN_LEFT = Cm(3)         # переплёт
PAGE_MARGIN_RIGHT = Cm(1.5)

TABLE_WIDTH_DXA = 9360           # ширина таблицы из оригинала
TABLE_BORDER_SIZE = 12           # 1.5pt в восьмых долях пункта

H1_TITLES = frozenset({"СОДЕРЖАНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ЛИТЕРАТУРЫ"})


# ---------------------------------------------------------------- утилиты

def fix_dashes(text: str | None) -> str:
    """Нормализация тире.

    1. Между цифрами — обычный дефис без пробелов: "2025-2026 гг.",
       "ст. 61.11-61.12" (это диапазон, а не пауза в речи).
    2. Все прочие длинные тире (—) — на короткие (–).

    Порядок важен: диапазоны обрабатываются первыми, иначе общая замена
    успеет превратить дефис в тире раньше времени.
    """
    if not text:
        return "" if text is None else str(text)
    result = str(text)
    result = re.sub(r"(\d)\s*[—–-]\s*(\d)", r"\1-\2", result)
    result = result.replace("—", "–")
    return result


def split_paragraphs(text: str | None) -> list[str]:
    """Разбивает текст на абзацы по пустым строкам."""
    if not text:
        return []
    parts = re.split(r"\n\s*\n", str(text))
    return [p.strip() for p in parts if p.strip()]


def strip_duplicate_heading(text: str | None, heading: str | None) -> str:
    """Убирает заголовок, если модель продублировала его первой строкой.

    Заголовок добавляется отдельно при вёрстке, поэтому "Введение" или
    "1.1. Название" в начале текста — это дубль.

    Срабатывает только на КОРОТКОЙ первой строке: длинная строка — уже
    основной текст, и отрезать её нельзя.
    """
    raw = str(text or "")
    if not heading:
        return raw

    lines = raw.split("\n")
    first_line = (lines[0] if lines else "").strip()
    # Порог с запасом 20 символов: заголовок может включать полное название.
    max_len = max(80, len(heading) + 20)
    if not first_line or len(first_line) > max_len:
        return raw

    def normalize(s: str) -> str:
        s = s.lower()
        s = re.sub(r"[«»\"'.:#*]", "", s)
        return re.sub(r"\s+", " ", s).strip()

    norm_first = normalize(first_line)
    norm_heading = normalize(heading)

    is_duplicate = (
        norm_first == norm_heading
        or (len(norm_heading) > 0 and norm_first.startswith(norm_heading))
        or norm_first in ("введение", "заключение")
    )
    if is_duplicate:
        return "\n".join(lines[1:]).strip()
    return raw


def parse_markdown_table(markdown: str | None) -> list[list[str]] | None:
    """Разбирает markdown-таблицу, которую выдала модель.

    Тонкость, ради которой парсер не наивный: модель иногда переносит
    длинный текст ячейки на следующую строку БЕЗ "|" в начале. Это не новая
    строка таблицы, а продолжение предыдущей. Парсер, отбрасывающий строки
    без "|", молча потеряет текст и сдвинет данные по столбцам — это
    незаметное искажение данных, а не косметический баг.
    """
    if not markdown:
        return None

    all_lines = str(markdown).split("\n")
    table_idx = [i for i, l in enumerate(all_lines) if l.strip().startswith("|")]
    if len(table_idx) < 2:
        return None

    zone = all_lines[table_idx[0]: table_idx[-1] + 1]

    # Склейка переносов внутри ячейки.
    merged: list[str] = []
    for line in zone:
        trimmed = line.strip()
        if trimmed.startswith("|"):
            merged.append(trimmed)
        elif trimmed and merged:
            merged[-1] += " " + trimmed

    def cells(line: str) -> list[str]:
        parts = [c.strip() for c in line.split("|")]
        if parts and parts[0] == "":
            parts = parts[1:]
        if parts and parts[-1] == "":
            parts = parts[:-1]
        return parts

    all_rows = [r for r in (cells(l) for l in merged) if r]
    # Строка-разделитель |---|---| состоит только из дефисов и двоеточий.
    rows = [r for r in all_rows if not all(re.fullmatch(r":?-+:?", c) for c in r)]
    if len(rows) < 2:
        return None

    # Нормализация числа столбцов по заголовку: лишние хвостовые ячейки
    # склеиваются в последнюю, а не обрезаются — иначе теряются данные.
    col_count = len(rows[0])
    normalized: list[list[str]] = []
    for r in rows:
        if len(r) <= col_count:
            normalized.append(r)
        else:
            head = r[: col_count - 1]
            tail = " | ".join(r[col_count - 1:])
            normalized.append([*head, tail])
    return normalized


# ------------------------------------------------------- построение блоков

def _style_run(run, size: Pt, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    # Кириллица в Word берёт шрифт из w:cs/w:eastAsia — без этого
    # Times New Roman может не примениться к русскому тексту.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def _zero_spacing(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING


def add_empty_paragraph(doc: Document):
    """Пустой абзац как явный визуальный отступ (вместо spacing)."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    _style_run(p.add_run(""), BODY_SIZE)
    return p


def add_h1(doc: Document, text: str, page_break_before: bool = True) -> None:
    """H1: КАПС, не жирный, по центру, разрыв страницы, + 2 пустых абзаца."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_break_before:
        p.add_run().add_break(WD_BREAK.PAGE)
    _style_run(p.add_run(fix_dashes(str(text)).upper()), H1_SIZE, bold=False)
    add_empty_paragraph(doc)
    add_empty_paragraph(doc)


def add_h2(doc: Document, text: str) -> None:
    """H2: обычный шрифт, без разрыва страницы."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    _style_run(p.add_run(fix_dashes(str(text))), H2_SIZE, bold=False)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Основной текст: 14pt, 1.5, по ширине, красная строка."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    _style_run(p.add_run(fix_dashes(text)), BODY_SIZE)


def add_body_paragraphs(doc: Document, text: str | None) -> None:
    for para in split_paragraphs(text):
        add_body_paragraph(doc, para)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(TABLE_BORDER_SIZE))  # 1.5pt
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def add_table(doc: Document, rows: Sequence[Sequence[str]] | None):
    """Таблица без заливки, 11pt, шапка по центру, тело по левому краю."""
    if not rows:
        return None

    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_width = Twips(TABLE_WIDTH_DXA // col_count)

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            if j >= col_count:
                break
            cell = cells[j]
            cell.width = col_width
            _set_cell_borders(cell)
            para = cell.paragraphs[0]
            para.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            _style_run(para.add_run(fix_dashes(value) or ""), TABLE_SIZE)
    return table


def add_table_caption(doc: Document, number: str, title: str) -> None:
    """Подпись "Таблица № N. Название" — отдельным абзацем перед таблицей."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    _style_run(p.add_run(fix_dashes(f"Таблица № {number}. {title}")), BODY_SIZE)


def add_contents_line(doc: Document, text: str, indented: bool = False) -> None:
    """Строка СОДЕРЖАНИЯ без номера страницы.

    Номера проставляются в Word после финальной вёрстки — до неё они
    неизвестны и любое проставленное число будет враньём.
    """
    p = doc.add_paragraph()
    _zero_spacing(p)
    if indented:
        p.paragraph_format.left_indent = CONTENTS_INDENT
    _style_run(p.add_run(fix_dashes(str(text))), BODY_SIZE)


def _iter_chapters(sections: Iterable[Mapping]):
    """Отдаёт (номер_главы|None, секция) — глава меняется по номеру до точки."""
    last_chapter = None
    for s in sections or []:
        number = str(s.get("number") or "")
        chapter_num = number.split(".")[0] if number else ""
        new_chapter = chapter_num if chapter_num and chapter_num != last_chapter else None
        if new_chapter:
            last_chapter = chapter_num
        yield new_chapter, s


def build_contents_page(
    doc: Document,
    sections: Sequence[Mapping] | None,
    chapter_titles: Mapping[str, str] | None,
    section_titles: Mapping[str, str] | None,
) -> None:
    add_h1(doc, "СОДЕРЖАНИЕ", page_break_before=False)
    add_contents_line(doc, "ВВЕДЕНИЕ")

    for new_chapter, s in _iter_chapters(sections or []):
        if new_chapter:
            name = (chapter_titles or {}).get(new_chapter)
            add_contents_line(
                doc, f"ГЛАВА {new_chapter}. {name}" if name else f"ГЛАВА {new_chapter}"
            )
        number = str(s.get("number") or "")
        if number:
            name = (section_titles or {}).get(number)
            add_contents_line(
                doc, f"{number}. {name}" if name else number, indented=True
            )

    add_contents_line(doc, "ЗАКЛЮЧЕНИЕ")


def _new_document() -> Document:
    """A4 + поля по ГОСТ + Times New Roman 14pt как стиль по умолчанию."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = PAGE_MARGIN_TOP
    section.bottom_margin = PAGE_MARGIN_BOTTOM
    section.left_margin = PAGE_MARGIN_LEFT
    section.right_margin = PAGE_MARGIN_RIGHT

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    return doc


def _add_section_blocks(
    doc: Document, heading: str | None, text: str | None, table_data
) -> None:
    """H2 + абзацы + опциональная таблица с подписью."""
    if heading:
        add_h2(doc, heading)
    add_body_paragraphs(doc, strip_duplicate_heading(text, heading))

    if not table_data:
        return
    # Совместимость: старый формат — просто markdown-строка без подписи.
    is_legacy = isinstance(table_data, str)
    markdown = table_data if is_legacy else table_data.get("markdown")
    rows = parse_markdown_table(markdown)
    if not rows:
        return
    if not is_legacy and table_data.get("number") and table_data.get("title"):
        add_table_caption(doc, table_data["number"], table_data["title"])
    add_table(doc, rows)


# ------------------------------------------------------------ точки входа

def generate_fragment_docx(
    *,
    title: str | None = None,
    text: str | None = None,
    table_markdown: str | None = None,
    is_h1: bool | None = None,
    chapter_heading: str | None = None,
    table_number: str | None = None,
    table_title: str | None = None,
    reference_sentence: str | None = None,
) -> bytes:
    """DOCX одного фрагмента (любой шаг протокола).

    Заголовок форматируется как H1, если это стандартный раздел работы
    или явно указан is_h1; иначе как H2 (например, "План раздела 1.1").
    """
    doc = _new_document()
    normalized_title = (title or "Документ").strip()
    use_h1 = is_h1 if is_h1 is not None else normalized_title.upper() in H1_TITLES

    if chapter_heading:
        add_h1(doc, chapter_heading, page_break_before=False)

    if use_h1:
        add_h1(doc, normalized_title, page_break_before=False)
    else:
        add_h2(doc, normalized_title)

    has_numbered_table = bool(table_markdown and table_number and table_title)
    body = strip_duplicate_heading(text, normalized_title)
    if has_numbered_table and reference_sentence:
        body = body.strip() + "\n\n" + reference_sentence
    add_body_paragraphs(doc, body)

    rows = parse_markdown_table(table_markdown)
    if rows:
        if has_numbered_table:
            add_table_caption(doc, table_number, table_title)
        add_table(doc, rows)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_full_docx(
    *,
    topic: str | None = None,
    introduction: str | None = None,
    sections: Sequence[Mapping] | None = None,
    conclusion: str | None = None,
    chapter_titles: Mapping[str, str] | None = None,
    section_titles: Mapping[str, str] | None = None,
) -> bytes:
    """DOCX всей работы: СОДЕРЖАНИЕ, ВВЕДЕНИЕ, главы с разделами, ЗАКЛЮЧЕНИЕ.

    Разделы группируются по главам на основе номера до точки
    (1.1, 1.2 → ГЛАВА 1; 2.1 → ГЛАВА 2).
    """
    doc = _new_document()

    build_contents_page(doc, sections, chapter_titles, section_titles)

    add_h1(doc, "ВВЕДЕНИЕ", page_break_before=True)
    add_body_paragraphs(doc, strip_duplicate_heading(introduction, "Введение"))

    for new_chapter, s in _iter_chapters(sections or []):
        if new_chapter:
            name = (chapter_titles or {}).get(new_chapter)
            heading = f"ГЛАВА {new_chapter}. {name}" if name else f"ГЛАВА {new_chapter}"
            add_h1(doc, heading, page_break_before=True)
        number = str(s.get("number") or "")
        name = (section_titles or {}).get(number) if number else None
        section_heading = f"{number}. {name}" if name else (number or None)
        _add_section_blocks(doc, section_heading, s.get("text"), s.get("table"))

    add_h1(doc, "ЗАКЛЮЧЕНИЕ", page_break_before=True)
    add_body_paragraphs(doc, strip_duplicate_heading(conclusion, "Заключение"))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
